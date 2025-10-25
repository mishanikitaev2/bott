import os
import logging
import re
from datetime import datetime
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, CallbackContext, 
    CallbackQueryHandler, MessageHandler, filters,
    ConversationHandler
)
from docx import Document
import tempfile
import shutil
from dotenv import load_dotenv

# Загружаем переменные из .env файла
load_dotenv()

# Настройка логирования
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# ==== ПОЛУЧАЕМ ПЕРЕМЕННЫЕ ИЗ .env ====
BOT_TOKEN = os.getenv('BOT_TOKEN')
ADMINS_STRING = os.getenv('ADMINS', '')  # Получаем строку с ID админов

# Преобразуем строку админов в список чисел
ADMINS = []
if ADMINS_STRING:
    try:
        ADMINS = [int(admin_id.strip()) for admin_id in ADMINS_STRING.split(',')]
    except ValueError as e:
        print(f"❌ Ошибка преобразования ADMINS: {e}")
        ADMINS = []

# Проверяем что токен и админы загружены
if not BOT_TOKEN:
    print("❌ BOT_TOKEN не найден в .env файле!")
    exit(1)

if not ADMINS:
    print("❌ ADMINS не найдены в .env файле!")
    exit(1)

print(f"✅ Токен загружен: {'*' * 10}{BOT_TOKEN[-5:]}")
print(f"✅ Админы: {ADMINS}")

# Состояния диалога
SELECTING_CATEGORY, SELECTING_TEMPLATES, FILLING_DATA = range(3)

# СТРУКТУРА КАТЕГОРИЙ И ШАБЛОНОВ
CATEGORIES = {
    "ОМС": {
        "ОМС": "ОМС.docx"
    },
    "ВМП": {
        "ВМП_выписка": "ВМП_выписка.docx",
        "ВМП_направление": "ВМП_направление.docx", 
        "ВМП_протокол": "ВМП_протокол.docx"
    },
    "ВМП в ОМС": {
        "ВМП_ОМС_выписка": "ВМП_ОМС_выписка.docx",
        "ВМП_ОМС_направление": "ВМП_ОМС_направление.docx",
        "ВМП_ОМС_протокол": "ВМП_ОМС_протокол.docx"
    }
}

def analyze_docx_template(template_path):
    """Анализирует .docx шаблон и возвращает список полей которые нужно заполнить"""
    try:
        if not os.path.exists(template_path):
            print(f"❌ Файл {template_path} не найден!")
            return []
        
        doc = Document(template_path)
        fields = []
        
        # Ищем поля в формате {field_name} во всех параграфах
        for paragraph in doc.paragraphs:
            found_fields = re.findall(r'\{(.*?)\}', paragraph.text)
            for field in found_fields:
                # ИСКЛЮЧАЕМ поля которые не нужно заполнять
                if field not in ['hist_number', 'current_date'] and field not in fields:
                    fields.append(field)
        
        # Ищем поля в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    found_fields = re.findall(r'\{(.*?)\}', cell.text)
                    for field in found_fields:
                        # ИСКЛЮЧАЕМ поля которые не нужно заполнять
                        if field not in ['hist_number', 'current_date'] and field not in fields:
                            fields.append(field)
        
        print(f"✅ В шаблоне {template_path} найдены поля: {fields}")
        return fields
        
    except Exception as e:
        print(f"❌ Ошибка анализа шаблона {template_path}: {e}")
        return []

def get_required_fields(selected_templates, category):
    """Возвращает все уникальные поля для выбранных шаблонов в ПОРЯДКЕ ИЗ ДОКУМЕНТОВ"""
    all_fields = []
    
    # Собираем поля из ВСЕХ шаблонов категории в порядке их появления
    for template_name in CATEGORIES[category].keys():
        template_file = CATEGORIES[category][template_name]
        template_path = f"templates/{template_file}"
        fields = analyze_docx_template(template_path)
        
        # Добавляем поля сохраняя порядок из документа и убирая дубли
        for field in fields:
            if field not in all_fields:
                all_fields.append(field)
    
    # Теперь оставляем только те поля, которые есть в ВЫБРАННЫХ шаблонах
    final_fields = []
    selected_fields_set = set()
    
    # Сначала собираем все поля из выбранных шаблонов
    for template_name in selected_templates:
        template_file = CATEGORIES[category][template_name]
        template_path = f"templates/{template_file}"
        fields = analyze_docx_template(template_path)
        selected_fields_set.update(fields)
    
    # Затем сохраняем порядок из all_fields, но только для выбранных полей
    for field in all_fields:
        if field in selected_fields_set:
            final_fields.append(field)
    
    # ОСОБЫЙ ПОРЯДОК: адрес проживания сразу после адреса регистрации
    if "address_fact" in final_fields and "address" in final_fields:
        address_index = final_fields.index("address")
        # Удаляем address_fact из текущей позиции
        final_fields.remove("address_fact")
        # Вставляем сразу после address
        final_fields.insert(address_index + 1, "address_fact")
    
    print(f"📋 Всего полей для заполнения: {len(final_fields)}")
    print(f"📋 Порядок полей: {final_fields}")
    return final_fields

def get_user_input_fields(required_fields):
    """Возвращает только те поля, которые действительно нужно спрашивать у пользователя"""
    # Поля, которые заполняются автоматически
    AUTO_FILLED_FIELDS = ["sop_diagnosis", "main_diagnosis"]
    
    user_fields = []
    for field in required_fields:
        if field not in AUTO_FILLED_FIELDS:
            user_fields.append(field)
    
    print(f"🎯 Поля для ввода пользователем: {len(user_fields)} из {len(required_fields)}")
    print(f"🎯 Список: {user_fields}")
    return user_fields

async def start(update: Update, context: CallbackContext):
    """Начало работы с ботом"""
    user_id = update.effective_user.id
    
    # Проверка доступа
    if user_id not in ADMINS:
        await update.message.reply_text("❌ Доступ запрещен.")
        return ConversationHandler.END
    
    # Очищаем предыдущие данные
    context.user_data.clear()
    
    # Создаем клавиатуру для выбора категории
    keyboard = []
    for category_name in CATEGORIES.keys():
        keyboard.append([InlineKeyboardButton(category_name, callback_data=f"category_{category_name}")])
    keyboard.append([InlineKeyboardButton("🔄 Перезапустить бота", callback_data="restart")])
    
    await update.message.reply_text(
        "🏥 Выбери тип медицинской помощи:\n\n"
        "• ОМС - базовый полис\n"
        "• ВМП - высокотехнологичная помощь\n"  
        "• ВМП в ОМС - ВМП по полису\n\n"
        "Выбери категорию:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    
    return SELECTING_CATEGORY

async def handle_category_selection(update: Update, context: CallbackContext):
    """Обработка выбора категории"""
    query = update.callback_query
    await query.answer()
    
    print(f"🔘 Нажата кнопка: {query.data}")
    
    if query.data == "restart":
        await query.edit_message_text("🔄 Перезапускаю бота...")
        context.user_data.clear()
        return await start_from_query(query, context)
    
    if query.data.startswith("category_"):
        category = query.data.replace("category_", "")
        context.user_data['category'] = category
        context.user_data['selected_templates'] = []  # Сбрасываем выбранные шаблоны
        
        # Создаем клавиатуру для выбора шаблонов в этой категории
        keyboard = []
        templates = CATEGORIES[category]
        
        for template_name in templates.keys():
            keyboard.append([InlineKeyboardButton(template_name, callback_data=template_name)])
        
        keyboard.append([InlineKeyboardButton("✅ Выбрать все", callback_data="select_all")])
        keyboard.append([InlineKeyboardButton("🚀 Продолжить", callback_data="continue")])
        keyboard.append([InlineKeyboardButton("◀️ Назад к категориям", callback_data="back_to_categories")])
        keyboard.append([InlineKeyboardButton("🔄 Перезапустить", callback_data="restart")])
        
        category_descriptions = {
            "ОМС": "📄 Базовые документы по полису ОМС",
            "ВМП": "🔬 Высокотехнологичная медицинская помощь", 
            "ВМП в ОМС": "💊 ВМП в рамках обязательного медицинского страхования"
        }
        
        await query.edit_message_text(
            f"{category_descriptions.get(category, category)}\n\n"
            "📋 Выбери нужные документы:\n\n"
            "• Нажми на названия которые нужны\n"
            "• Они выделятся галочкой\n"  
            "• Можно выбрать все сразу или по отдельности\n"
            "• Когда выбрал нужные - жми '🚀 Продолжить'",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        
        return SELECTING_TEMPLATES
    
    return SELECTING_CATEGORY

async def handle_template_selection(update: Update, context: CallbackContext):
    """Обработка выбора шаблонов - ИСПРАВЛЕННАЯ ВЕРСИЯ БЕЗ MARKDOWN"""
    query = update.callback_query
    await query.answer()
    
    print(f"🔘 Нажата кнопка: {query.data}")
    
    if query.data == "back_to_categories":
        # Возвращаемся к выбору категории
        keyboard = []
        for category_name in CATEGORIES.keys():
            keyboard.append([InlineKeyboardButton(category_name, callback_data=f"category_{category_name}")])
        keyboard.append([InlineKeyboardButton("🔄 Перезапустить бота", callback_data="restart")])
        
        await query.edit_message_text(
            "🏥 Выбери тип медицинской помощи:\n\n"
            "• ОМС - базовый полис\n"
            "• ВМП - высокотехнологичная помощь\n"  
            "• ВМП в ОМС - ВМП по полису\n\n"
            "Выбери категорию:",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return SELECTING_CATEGORY
    
    if query.data == "restart":
        await query.edit_message_text("🔄 Перезапускаю бота...")
        context.user_data.clear()
        return await start_from_query(query, context)
    
    if query.data == "select_all":
        category = context.user_data.get('category')
        if category:
            selected = list(CATEGORIES[category].keys())
            context.user_data['selected_templates'] = selected
            
            # Обновляем сообщение с выбранными документами
            keyboard = []
            templates = CATEGORIES[category]
            
            for template_name in templates.keys():
                keyboard.append([InlineKeyboardButton(f"✅ {template_name}", callback_data=template_name)])
            
            keyboard.append([InlineKeyboardButton("✅ Выбрать все", callback_data="select_all")])
            keyboard.append([InlineKeyboardButton("🚀 Продолжить", callback_data="continue")])
            keyboard.append([InlineKeyboardButton("◀️ Назад к категориям", callback_data="back_to_categories")])
            keyboard.append([InlineKeyboardButton("🔄 Перезапустить", callback_data="restart")])
            
            await query.edit_message_text(
                f"✅ Выбраны ВСЕ документы для {category}:\n"
                f"📝 {', '.join(selected)}\n\n"
                f"Нажми '🚀 Продолжить' для заполнения данных",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
            return SELECTING_TEMPLATES
    
    if query.data == "continue":
        selected = context.user_data.get('selected_templates', [])
        category = context.user_data.get('category')
        
        print(f"🎯 Кнопка 'Продолжить' нажата. Выбрано: {selected}, Категория: {category}")
        
        if not selected:
            await query.answer("❌ Нужно выбрать хотя бы один документ!", show_alert=True)
            return SELECTING_TEMPLATES
        
        # ИСПРАВЛЕНИЕ: используем reply_text вместо edit_message_text для нового сообщения
        await query.edit_message_text(
            f"✅ Выбрано документов: {len(selected)}\n"
            f"📝 {', '.join(selected)}\n\n"
            f"🔄 Переходим к заполнению данных..."
        )
        
        # Анализируем какие поля нужны для выбранных шаблонов
        required_fields = get_required_fields(selected, category)
        
        if not required_fields:
            await context.bot.send_message(
                query.message.chat.id,
                "❌ В выбранных шаблонах нет полей для заполнения!\n"
                "Проверь что файлы шаблонов существуют.\n\n"
                "Нажми /start чтобы начать заново."
            )
            return ConversationHandler.END
        
        # Получаем только те поля, которые действительно нужно спрашивать у пользователя
        user_input_fields = get_user_input_fields(required_fields)
        
        # Сохраняем оба списка
        context.user_data['required_fields'] = required_fields  # Все поля для документов
        context.user_data['user_input_fields'] = user_input_fields  # Только для вопросов
        context.user_data['current_field_index'] = 0
        context.user_data['field_history'] = []  # История заполненных полей для отмены
        
        # Начинаем заполнение
        await ask_next_question(context, query.message.chat.id)
        return FILLING_DATA
    
    # Основная логика выбора/отмены шаблонов
    category = context.user_data.get('category')
    if not category:
        await query.edit_message_text("❌ Ошибка: категория не выбрана. Нажми /start")
        return ConversationHandler.END
    
    selected = context.user_data.get('selected_templates', [])
    template_name = query.data
    
    if template_name in selected:
        selected.remove(template_name)
        print(f"➖ Убрали шаблон: {template_name}")
    else:
        selected.append(template_name)
        print(f"➕ Добавили шаблон: {template_name}")
    
    context.user_data['selected_templates'] = selected
    
    # Обновляем клавиатуру с отметками
    keyboard = []
    templates = CATEGORIES[category]
    
    for template_name in templates.keys():
        emoji = "✅" if template_name in selected else "◻️"
        keyboard.append([InlineKeyboardButton(f"{emoji} {template_name}", callback_data=template_name)])
    
    keyboard.append([InlineKeyboardButton("✅ Выбрать все", callback_data="select_all")])
    keyboard.append([InlineKeyboardButton("🚀 Продолжить", callback_data="continue")])
    keyboard.append([InlineKeyboardButton("◀️ Назад к категориям", callback_data="back_to_categories")])
    keyboard.append([InlineKeyboardButton("🔄 Перезапустить", callback_data="restart")])
    
    await query.edit_message_text(
        f"📋 Выбери нужные документы:\n\n"
        f"• Выбрано: {len(selected)}/{len(templates)}\n"
        f"• Нажми на названия которые нужны\n" 
        f"• Они выделятся галочкой\n"
        f"• Можно выбрать все сразу или по отдельности\n"
        f"• Когда выбрал нужные - жми '🚀 Продолжить'",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    
    return SELECTING_TEMPLATES

async def start_from_query(query, context):
    """Запуск бота из callback query"""
    keyboard = []
    for category_name in CATEGORIES.keys():
        keyboard.append([InlineKeyboardButton(category_name, callback_data=f"category_{category_name}")])
    keyboard.append([InlineKeyboardButton("🔄 Перезапустить бота", callback_data="restart")])
    
    await query.message.reply_text(
        "🏥 Выбери тип медицинской помощи:\n\n"
        "• ОМС - базовый полис\n"
        "• ВМП - высокотехнологичная помощь\n"  
        "• ВМП в ОМС - ВМП по полису\n\n"
        "Выбери категорию:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    
    return SELECTING_CATEGORY

async def ask_next_question(context: CallbackContext, chat_id: int):
    """Задаем следующий вопрос - ИСПРАВЛЕНО: используем только user_input_fields"""
    field_index = context.user_data['current_field_index']
    user_input_fields = context.user_data['user_input_fields']
    
    # ВАЖНОЕ ИСПРАВЛЕНИЕ: проверяем границы массива
    if field_index >= len(user_input_fields):
        print("✅ Все поля заполнены, переходим к генерации документов")
        await generate_documents(context, chat_id)
        return ConversationHandler.END
    
    print(f"📝 Заполняем поле {field_index + 1}/{len(user_input_fields)}: {user_input_fields[field_index]}")
    
    field_name = user_input_fields[field_index]
    
    # Названия полей как в документах
    field_display_names = {
        # Основные данные
        "name": "👤 ФИО пациента",
        "birth_date": "📅 Дата рождения (ДД.ММ.ГГГГ)",
        "address": "📍 Адрес регистрации по месту жительства",
        "address_fact": "🏠 Адрес фактического проживания",
        
        # Документы
        "oms": "📋 Номер полиса ОМС",
        "snils": "📘 СНИЛС",
        
        # Диагнозы
        "diagnosis": "🏥 Установлен клинический диагноз",
        "diagnosis_code": "🔢 Код по МКБ-10",
        
        # Медицинская информация
        "medical_history": "📋 Anamnesis morbi",
        "status_localis": "📊 Status localis",
        
        # ВМП данные
        "wmp": "🔬 Наименование вида ВМП",
        "wmp_oms": "💊 Наименование вида ВМП в ОМС",
        "wmp_group": "📁 № группы ВМП",
        "wmp_code": "🔢 Код вида ВМП",
        "wmp_oms_group": "📂 № группы ВМП в ОМС", 
        "wmp_oms_code": "🔣 Код вида ВМП в ОМС",
        "patient_model": "👥 Модель пациента",
        "treatment_method": "💉 Метод лечения ВМП",
        
        # ОМС данные
        "ksg_group": "📊 Группа КСГ",
        "operation_code": "🔪 Код операции",
        
        # Заключения
        "recommendations": "📝 Рекомендации / Решение комиссии",
        
        # Врачи
        "doctor": "👨‍⚕️ ФИО врача",
        "fio_lech": "👩‍⚕️ ФИО лечащего врача (для подписи)",
        "department": "🏢 Отделение"
    }
    
    question = field_display_names.get(field_name, f"📝 {field_name}")
    
    # Добавляем прогресс-бар
    progress = f"({field_index + 1}/{len(user_input_fields)})"
    
    # Добавляем кнопки навигации с возможностью отмены предыдущего шага
    keyboard = []
    if field_index > 0:  # Если не первое поле - показываем кнопку "Назад"
        keyboard.append([InlineKeyboardButton("◀️ Исправить предыдущее поле", callback_data="back_to_previous")])
    
    keyboard.append([
        InlineKeyboardButton("◀️ Назад к выбору", callback_data="back_to_templates"),
        InlineKeyboardButton("🔄 Перезапустить", callback_data="restart")
    ])
    
    await context.bot.send_message(
        chat_id, 
        f"{progress} {question}",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def handle_user_input(update: Update, context: CallbackContext):
    """Обработка ввода пользователя - ИЗМЕНЕНО: все диагнозы = клинический диагноз"""
    user_input = update.message.text
    field_index = context.user_data['current_field_index']
    user_input_fields = context.user_data['user_input_fields']
    
    # Сохраняем данные
    field_name = user_input_fields[field_index]
    context.user_data[field_name] = user_input
    
    # НОВАЯ ЛОГИКА: если заполняем "diagnosis" (клинический диагноз), 
    # то автоматически заполняем ВСЕ связанные диагнозы тем же значением
    if field_name == "diagnosis":
        # Автоматически заполняем все связанные поля диагнозов
        context.user_data["sop_diagnosis"] = user_input  # сопутствующий
        context.user_data["main_diagnosis"] = user_input  # основной
        print(f"💡 Автоматически заполнены все диагнозы: {user_input}")
    
    # Сохраняем в историю для возможности отмены
    if 'field_history' not in context.user_data:
        context.user_data['field_history'] = []
    
    context.user_data['field_history'].append({
        'field_name': field_name,
        'value': user_input,
        'index': field_index
    })
    
    print(f"💾 Сохранено поле {field_name}: {user_input}")
    
    # Переходим к следующему полю
    context.user_data['current_field_index'] += 1
    await ask_next_question(context, update.effective_chat.id)
    
    return FILLING_DATA

async def handle_navigation(update: Update, context: CallbackContext):
    """Обработка навигационных кнопок"""
    query = update.callback_query
    await query.answer()
    
    print(f"🔘 Нажата навигационная кнопка: {query.data}")
    
    if query.data == "back_to_previous":
        # Возвращаемся к предыдущему полю для исправления
        current_index = context.user_data.get('current_field_index', 0)
        if current_index > 0:
            context.user_data['current_field_index'] = current_index - 1
            
            # Удаляем последнее значение из истории
            if context.user_data.get('field_history'):
                context.user_data['field_history'].pop()
            
            await query.edit_message_text("↩️ Возвращаюсь к предыдущему полю для исправления...")
            await ask_next_question(context, query.message.chat.id)
            return FILLING_DATA
        else:
            await query.answer("❌ Это первое поле, нельзя вернуться назад", show_alert=True)
            return FILLING_DATA
    
    elif query.data == "back_to_templates":
        category = context.user_data.get('category')
        if category:
            # Возвращаемся к выбору шаблонов
            keyboard = []
            templates = CATEGORIES[category]
            selected = context.user_data.get('selected_templates', [])
            
            for template_name in templates.keys():
                emoji = "✅" if template_name in selected else "◻️"
                keyboard.append([InlineKeyboardButton(f"{emoji} {template_name}", callback_data=template_name)])
            
            keyboard.append([InlineKeyboardButton("✅ Выбрать все", callback_data="select_all")])
            keyboard.append([InlineKeyboardButton("🚀 Продолжить", callback_data="continue")])
            keyboard.append([InlineKeyboardButton("◀️ Назад к категориям", callback_data="back_to_categories")])
            keyboard.append([InlineKeyboardButton("🔄 Перезапустить", callback_data="restart")])
            
            await query.edit_message_text(
                f"📋 Выбери нужные документы:\n\n"
                f"• Выбрано: {len(selected)}/{len(templates)}\n"
                f"• Нажми на названия которые нужны\n" 
                f"• Они выделятся галочкой\n"
                f"• Можно выбрать все сразу или по отдельности\n"
                f"• Когда выбрал нужные - жми '🚀 Продолжить'",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
            return SELECTING_TEMPLATES
    
    elif query.data == "restart":
        await query.edit_message_text("🔄 Перезапускаю бота...")
        context.user_data.clear()
        return await start_from_query(query, context)
    
    return FILLING_DATA

def replace_in_paragraph(paragraph, data):
    """Заменяет плейсхолдеры в параграфе с сохранением форматирования"""
    for key, value in data.items():
        placeholder = f"{{{key}}}"
        if placeholder in paragraph.text:
            # Сохраняем стиль
            original_runs = paragraph.runs
            if original_runs:
                # Сохраняем свойства первого run
                first_run = original_runs[0]
                original_bold = first_run.bold
                original_italic = first_run.italic
                original_underline = first_run.underline
                
                # Сохраняем шрифт если возможно
                try:
                    original_font_name = first_run.font.name
                    original_font_size = first_run.font.size
                except:
                    original_font_name = None
                    original_font_size = None
            else:
                original_bold = None
                original_italic = None
                original_underline = None
                original_font_name = None
                original_font_size = None
            
            # Заменяем текст
            paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Восстанавливаем стиль
            if paragraph.runs:
                run = paragraph.runs[0]
                if original_bold is not None:
                    run.bold = original_bold
                if original_italic is not None:
                    run.italic = original_italic
                if original_underline is not None:
                    run.underline = original_underline
                
                # Восстанавливаем шрифт
                try:
                    if original_font_name:
                        run.font.name = original_font_name
                    if original_font_size:
                        run.font.size = original_font_size
                except:
                    pass  # Игнорируем ошибки шрифта

def fill_docx_template(template_path, data):
    """Заполняет .docx шаблон данными с сохранением форматирования"""
    try:
        # Загружаем оригинальный шаблон
        doc = Document(template_path)
        
        # Заполняем плейсхолдеры в параграфах
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, data)
        
        # Заполняем плейсхолдеры в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, data)
        
        return doc
        
    except Exception as e:
        print(f"❌ Ошибка заполнения шаблона {template_path}: {e}")
        # Создаем простой документ в случае ошибки
        doc = Document()
        doc.add_heading('МЕДИЦИНСКИЙ ДОКУМЕНТ', 0)
        for key, value in data.items():
            doc.add_paragraph(f"{key}: {value}")
        return doc

async def generate_documents(context: CallbackContext, chat_id: int):
    """Генерация и отправка Word документов"""
    user_data = context.user_data
    selected_templates = user_data.get('selected_templates', [])
    category = user_data.get('category')
    
    if not selected_templates or not category:
        await context.bot.send_message(chat_id, "❌ Ошибка: не выбраны шаблоны или категория")
        return ConversationHandler.END
    
    # Собираем все данные для документов
    data = {}
    required_fields = user_data.get('required_fields', [])
    for field in required_fields:
        data[field] = user_data.get(field, "Не указано")
    
    # НОВАЯ ЛОГИКА: убедимся что ВСЕ диагнозы совпадают с клиническим
    if "diagnosis" in data:
        # Автоматически заполняем все связанные диагнозы
        data["sop_diagnosis"] = data["diagnosis"]  # сопутствующий
        data["main_diagnosis"] = data["diagnosis"]  # основной
        print(f"💡 Все диагнозы установлены равными клиническому: {data['diagnosis']}")
    
    print(f"🎯 Генерируем документы для {category}: {selected_templates}")
    
    temp_dir = tempfile.mkdtemp()
    
    try:
        generated_files = []
        
        for template_name in selected_templates:
            template_file = CATEGORIES[category][template_name]
            template_path = f"templates/{template_file}"
            
            safe_template_name = re.sub(r'[^\w\s-]', '', template_name)
            safe_template_name = safe_template_name.replace(' ', '_')
            
            if not os.path.exists(template_path):
                doc = Document()
                doc.add_heading(template_name, 0)
                for key, value in data.items():
                    doc.add_paragraph(f"{key}: {value}")
            else:
                doc = fill_docx_template(template_path, data)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_template_name}_{timestamp}.docx"
            file_path = os.path.join(temp_dir, filename)
            doc.save(file_path)
            generated_files.append((template_name, file_path))
            print(f"✅ Создан файл: {file_path}")
        
        if generated_files:
            await context.bot.send_message(chat_id, "📄 Генерирую документы...")
            
            for template_name, file_path in generated_files:
                safe_display_name = re.sub(r'[^\w\s-]', '', template_name)
                
                with open(file_path, 'rb') as doc_file:
                    await context.bot.send_document(
                        chat_id=chat_id,
                        document=doc_file,
                        filename=f"{safe_display_name}.docx",
                        caption=f"✅ {template_name}"
                    )
                print(f"📤 Отправлен файл: {file_path}")
            
            shutil.rmtree(temp_dir)
            print("🧹 Временные файлы удалены")
            
            keyboard = [
                [InlineKeyboardButton("🔄 Новый документ", callback_data="restart")]
            ]
            
            await context.bot.send_message(
                chat_id,
                "🎉 Все документы готовы!\n\n"
                "⚠️ Временные файлы удалены из системы\n\n"
                "Для нового документа нажми кнопку ниже:",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
        else:
            await context.bot.send_message(chat_id, "❌ Не удалось сгенерировать документы")
        
    except Exception as e:
        logger.error(f"Ошибка генерации: {e}")
        await context.bot.send_message(chat_id, f"❌ Произошла ошибка при генерации документов: {str(e)}")
        
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
    
    finally:
        context.user_data.clear()

async def cancel(update: Update, context: CallbackContext):
    """Отмена операции"""
    context.user_data.clear()
    await update.message.reply_text(
        "❌ Операция отменена.\n"
        "Все временные данные удалены.\n\n"
        "Для начала используй /start"
    )
    return ConversationHandler.END

def main():
    # Проверяем что переменные загружены
    if not BOT_TOKEN or not ADMINS:
        print("❌ Не удалось загрузить переменные из .env файла!")
        return
    
    print("🤖 Запускаю бота...")
    print("🔍 Проверяю шаблоны...")
    
    # Проверяем все шаблоны
    for category, templates in CATEGORIES.items():
        print(f"\n📁 Категория: {category}")
        for template_name, template_file in templates.items():
            template_path = f"templates/{template_file}"
            if os.path.exists(template_path):
                fields = analyze_docx_template(template_path)
                print(f"   ✅ {template_name}: {len(fields)} полей")
            else:
                print(f"   ❌ {template_name}: файл не найден")
    
    application = Application.builder().token(BOT_TOKEN).build()
    
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler('start', start)],
        states={
            SELECTING_CATEGORY: [
                CallbackQueryHandler(handle_category_selection)
            ],
            SELECTING_TEMPLATES: [
                CallbackQueryHandler(handle_template_selection)
            ],
            FILLING_DATA: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_user_input),
                CallbackQueryHandler(handle_navigation)
            ]
        },
        fallbacks=[
            CommandHandler('cancel', cancel),
            CommandHandler('start', start)
        ]
    )
    
    application.add_handler(conv_handler)
    application.add_handler(CommandHandler("cancel", cancel))
    application.add_handler(CommandHandler("start", start))
    
    print("\n✅ Бот запущен!")
    print("📱 Телеграм -> /start")
    print("⏹️  Ctrl+C для остановки")
    
    application.run_polling()

if __name__ == '__main__':
    main()
